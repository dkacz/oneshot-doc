"""Post-process a Quarto/Pandoc DOCX with the oneshot-doc house layout.

Run after every `quarto render ... --to docx`. The script is intentionally
generic and idempotent: repeated runs update the same layout surfaces instead
of appending duplicate headers, footers, labels, or address blocks.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Mm, Pt, RGBColor

try:
    from palette import get_theme, theme_names
except ImportError:  # pragma: no cover - only used when copied away from assets/
    def theme_names() -> list[str]:
        return ["think-tank"]

    def get_theme(name: str = "think-tank", accent: str | None = None) -> dict[str, str]:
        if name != "think-tank":
            raise ValueError("Unknown theme {!r}. Available themes: think-tank".format(name))
        theme = {
            "primary": "#253B59",
            "rule": "#7D93AD",
            "gray": "#4B5563",
            "light": "#EEF3F8",
            "accent": "#A6442C",
            "body_family": "serif",
            "heading_family": "sans",
        }
        if accent is not None:
            theme["accent"] = accent if accent.startswith("#") else "#" + accent
        return theme

SANS = "Source Sans 3"
SERIF = "Source Serif 4"
MONO = "Consolas"
DEFAULT_ADDRESS_COLS_TWIPS = (5200, 3870)
DEFAULT_MARGINS_MM = {"top": 20.0, "bottom": 20.0, "left": 24.0, "right": 24.0}
BOX_STYLES = {"KeyBox", "MethodBox"}
LIST_LIKE_STYLES = {"Compact", "List Paragraph", "List Bullet", "List Number"}


def theme_box_config(theme_name: str, theme: dict[str, str]) -> dict[str, dict[str, Any]]:
    key_accent = theme["primary"] if theme_name == "think-tank" else theme["accent"]
    return {
        "KeyBox": {
            "fill": theme["light"].lstrip("#"),
            "accent": key_accent.lstrip("#"),
            "border": theme["light"].lstrip("#"),
            "left_sz": "20",
            "border_sz": "2",
            "margins": {"top": "120", "left": "190", "bottom": "120", "right": "170"},
        },
        "MethodBox": {
            "fill": "FFFFFF",
            "accent": theme["rule"].lstrip("#"),
            "border": theme["rule"].lstrip("#"),
            "left_sz": "8",
            "border_sz": "4",
            "margins": {"top": "100", "left": "180", "bottom": "100", "right": "170"},
        },
    }


def rgb(hex_color: str) -> RGBColor:
    value = hex_color.lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def set_run_font(run, name: str, size: float | None = None, color: RGBColor | None = None) -> None:
    run.font.name = name
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "cs"):
        rfonts.set(qn("w:" + attr), name)


def set_style_font(style, name: str, size: float | None = None, color: RGBColor | None = None) -> None:
    style.font.name = name
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "cs"):
        rfonts.set(qn("w:" + attr), name)


def clear_paragraph(paragraph) -> None:
    ppr = paragraph._p.pPr
    for child in list(paragraph._p):
        if child is not ppr:
            paragraph._p.remove(child)
    ppr = paragraph._p.get_or_add_pPr()
    old_border = ppr.find(qn("w:pBdr"))
    if old_border is not None:
        ppr.remove(old_border)


def first_child(parent, tag: str):
    node = parent.find(qn(tag))
    if node is None:
        node = OxmlElement(tag)
        parent.append(node)
    return node


def remove_child(parent, tag: str) -> None:
    child = parent.find(qn(tag))
    if child is not None:
        parent.remove(child)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for element in (fld_begin, instr, fld_sep, text, fld_end):
        run._r.append(element)
    return run


def set_bottom_border(paragraph, color: str, size: str = "4") -> None:
    ppr = paragraph._p.get_or_add_pPr()
    old = ppr.find(qn("w:pBdr"))
    if old is not None:
        ppr.remove(old)
    ppr.append(
        parse_xml(
            f"<w:pBdr {nsdecls('w')}>"
            f'<w:bottom w:val="single" w:sz="{size}" w:space="4" w:color="{color}"/>'
            "</w:pBdr>"
        )
    )


def parse_geometry(value: str) -> tuple[str, dict[str, float]]:
    """Parse `A4,top=20,bottom=20,left=24,right=24` style geometry."""
    if ":" in value:
        value = value.replace(":", ",", 1)
    parts = [part.strip() for part in value.split(",") if part.strip()]
    page = parts[0].upper() if parts else "A4"
    if page != "A4":
        raise ValueError("Only A4 geometry is supported")
    margins = dict(DEFAULT_MARGINS_MM)
    for part in parts[1:]:
        if "=" not in part:
            continue
        key, raw = [x.strip().lower() for x in part.split("=", 1)]
        key = {"t": "top", "b": "bottom", "l": "left", "r": "right"}.get(key, key)
        if key not in margins and key not in {"header", "footer"}:
            raise ValueError(f"Unknown geometry key: {key}")
        margins[key] = float(raw.removesuffix("mm"))
    return page, margins


def apply_geometry(doc, geometry: str) -> None:
    _, margins = parse_geometry(geometry)
    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(margins["top"])
        section.bottom_margin = Mm(margins["bottom"])
        section.left_margin = Mm(margins["left"])
        section.right_margin = Mm(margins["right"])
        section.header_distance = Mm(margins.get("header", 10.0))
        section.footer_distance = Mm(margins.get("footer", 10.0))


def set_running_header(doc, text: str, theme_name: str, theme: dict[str, str]) -> None:
    for section in doc.sections:
        section.different_first_page_header_footer = True

        section.first_page_header.is_linked_to_previous = False
        clear_paragraph(section.first_page_header.paragraphs[0])

        header = section.header
        header.is_linked_to_previous = False
        paragraph = header.paragraphs[0]
        clear_paragraph(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(text)
        set_run_font(run, SANS, 9, rgb(theme["gray"]))
        set_bottom_border(paragraph, theme["rule"].lstrip("#"), "2" if theme_name == "minimal" else "4")


def set_page_footer(doc, text: str | None, theme: dict[str, str]) -> None:
    for section in doc.sections:
        section.different_first_page_header_footer = True
        body_width = section.page_width - section.left_margin - section.right_margin
        for footer in (section.first_page_footer, section.footer):
            footer.is_linked_to_previous = False
            paragraph = footer.paragraphs[0]
            clear_paragraph(paragraph)
            paragraph.paragraph_format.tab_stops.clear_all()
            if text:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.tab_stops.add_tab_stop(body_width, WD_TAB_ALIGNMENT.RIGHT)
                set_run_font(paragraph.add_run(text), SANS, 9, rgb(theme["gray"]))
                paragraph.add_run("\t")
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            page_run = add_page_field(paragraph)
            set_run_font(page_run, SANS, 9, rgb(theme["gray"]))


def ensure_doc_kind_style(doc, theme_name: str, theme: dict[str, str]):
    styles = doc.styles
    try:
        style = styles["DocKind"]
    except KeyError:
        style = styles.add_style("DocKind", WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles["Normal"]
    style.quick_style = True
    set_style_font(style, SANS, 8.5, rgb(theme["gray"]))
    style.font.small_caps = True
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(4)
    # The academic theme centers the kind label in the PDF title block.
    style.paragraph_format.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if theme_name == "academic" else WD_ALIGN_PARAGRAPH.LEFT
    )
    return style


def add_doc_label(doc, label: str, theme_name: str, theme: dict[str, str]) -> None:
    if not label.strip():
        return
    style = ensure_doc_kind_style(doc, theme_name, theme)
    paragraphs = list(doc.paragraphs)
    title = next((p for p in paragraphs if p.style.name == "Title"), None)
    if title is None:
        first = next((p for p in paragraphs if p.text.strip()), None)
        if first is not None and first.text.strip() == label:
            first.style = style
        return
    idx = paragraphs.index(title)
    if idx > 0 and paragraphs[idx - 1].text.strip() == label:
        paragraphs[idx - 1].style = style
        return
    existing = next((p for p in paragraphs if p.style.name == "DocKind" and p.text.strip() == label), None)
    if existing is not None:
        existing.style = style
        return
    paragraph = title.insert_paragraph_before(label)
    paragraph.style = style
    for run in paragraph.runs:
        set_run_font(run, SANS, 8.5, rgb(theme["gray"]))
        run.font.small_caps = True


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    old = tbl_pr.find(qn("w:tblBorders"))
    if old is not None:
        tbl_pr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "nil")
        borders.append(node)
    tbl_pr.append(borders)


def set_table_cell_margins(table, value: str = "0") -> None:
    tbl_pr = table._tbl.tblPr
    old = tbl_pr.find(qn("w:tblCellMar"))
    if old is not None:
        tbl_pr.remove(old)
    margins = OxmlElement("w:tblCellMar")
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), value)
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    tbl_pr.append(margins)


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.insert(0, tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, value: str = "0") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    old = tc_pr.find(qn("w:tcMar"))
    if old is not None:
        tc_pr.remove(old)
    margins = OxmlElement("w:tcMar")
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), value)
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    tc_pr.append(margins)


def set_cell_shading(cell, fill: str = "FFFFFF") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    old = tc_pr.find(qn("w:shd"))
    if old is not None:
        tc_pr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def paragraph_style_name(paragraph) -> str:
    try:
        return paragraph.style.name
    except Exception:  # pragma: no cover - defensive against damaged style refs
        return ""


def paragraph_has_numbering(paragraph) -> bool:
    ppr = paragraph._p.pPr
    return ppr is not None and ppr.find(qn("w:numPr")) is not None


def paragraph_is_list_like(paragraph) -> bool:
    return paragraph_has_numbering(paragraph) or paragraph_style_name(paragraph) in LIST_LIKE_STYLES


def remove_direct_paragraph_box_format(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    remove_child(ppr, "w:pBdr")
    remove_child(ppr, "w:shd")


def demote_box_styles_to_markers(doc) -> None:
    """Remove paragraph-frame rendering from box marker styles.

    Pandoc applies custom-style divs at paragraph level. If the style itself
    owns borders or shading, Word prints list bullets outside the visual box.
    The table wrapper below becomes the only visual frame.
    """

    for name in BOX_STYLES:
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        ppr = style.element.get_or_add_pPr()
        remove_child(ppr, "w:pBdr")
        remove_child(ppr, "w:shd")


def find_box_groups(doc) -> list[tuple[str, list[Any]]]:
    groups: list[tuple[str, list[Any]]] = []
    paragraphs = list(doc.paragraphs)
    i = 0
    while i < len(paragraphs):
        style = paragraph_style_name(paragraphs[i])
        if style not in BOX_STYLES:
            i += 1
            continue
        box_style = style
        group = [paragraphs[i]]
        i += 1
        while i < len(paragraphs):
            current_style = paragraph_style_name(paragraphs[i])
            if current_style == box_style or paragraph_is_list_like(paragraphs[i]):
                group.append(paragraphs[i])
                i += 1
                continue
            break
        groups.append((box_style, group))
    return groups


def insert_one_cell_table_before(doc, reference):
    table = doc.add_table(rows=1, cols=1)
    body = doc._body._body
    body.remove(table._tbl)
    if reference is not None:
        reference._p.addprevious(table._tbl)
    else:
        body.append(table._tbl)
    return table


def set_box_cell_margins(cell, margins: dict[str, str]) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    old = tc_pr.find(qn("w:tcMar"))
    if old is not None:
        tc_pr.remove(old)
    tc_mar = OxmlElement("w:tcMar")
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), margins[edge])
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def set_table_width_pct(table, pct: str = "5000") -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = first_child(tbl_pr, "w:tblW")
    tbl_w.set(qn("w:w"), pct)
    tbl_w.set(qn("w:type"), "pct")
    jc = first_child(tbl_pr, "w:jc")
    jc.set(qn("w:val"), "left")
    tbl_ind = first_child(tbl_pr, "w:tblInd")
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = first_child(tbl_pr, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")

    cell = table.cell(0, 0)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = first_child(tc_pr, "w:tcW")
    tc_w.set(qn("w:w"), pct)
    tc_w.set(qn("w:type"), "pct")


def set_box_table_borders(table, cfg: dict[str, Any]) -> None:
    tbl_pr = table._tbl.tblPr
    remove_child(tbl_pr, "w:tblBorders")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        if edge.startswith("inside"):
            node.set(qn("w:val"), "nil")
        else:
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), cfg["left_sz"] if edge == "left" else cfg["border_sz"])
            node.set(qn("w:space"), "0")
            node.set(qn("w:color"), cfg["accent"] if edge == "left" else cfg["border"])
        borders.append(node)
    tbl_pr.append(borders)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def clear_cell_paragraphs(cell) -> None:
    for paragraph in list(cell._tc.findall(qn("w:p"))):
        cell._tc.remove(paragraph)


def apply_box_table_geometry(table, cfg: dict[str, Any]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width_pct(table)
    set_box_table_borders(table, cfg)
    row = table.rows[0]
    set_row_cant_split(row)
    cell = row.cells[0]
    set_box_cell_margins(cell, cfg["margins"])
    set_cell_shading(cell, cfg["fill"])
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    clear_cell_paragraphs(cell)


def convert_boxes_to_tables(doc, theme_name: str, theme: dict[str, str]) -> int:
    demote_box_styles_to_markers(doc)
    config = theme_box_config(theme_name, theme)
    converted = 0
    for box_style, paragraphs in find_box_groups(doc):
        table = insert_one_cell_table_before(doc, paragraphs[0])
        apply_box_table_geometry(table, config[box_style])
        cell = table.cell(0, 0)
        for idx, paragraph in enumerate(paragraphs):
            remove_direct_paragraph_box_format(paragraph)
            paragraph.paragraph_format.keep_together = True
            if idx < len(paragraphs) - 1:
                paragraph.paragraph_format.keep_with_next = True
            cell._tc.append(paragraph._p)
        converted += 1
    return converted


def remove_row(row) -> None:
    row._element.getparent().remove(row._element)


def twips_to_mm(width_twips: int) -> float:
    return width_twips / 1440 * 25.4


def set_address_table_geometry(table, columns_twips: tuple[int, int]) -> None:
    total = sum(columns_twips)
    tbl_pr = table._tbl.tblPr
    tbl_w = first_child(tbl_pr, "w:tblW")
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    jc = first_child(tbl_pr, "w:jc")
    jc.set(qn("w:val"), "left")
    layout = first_child(tbl_pr, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_ind = first_child(tbl_pr, "w:tblInd")
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    set_table_cell_margins(table)

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(1, grid)
    for child in list(grid):
        grid.remove(child)
    for width in columns_twips:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells[: len(columns_twips)]):
            width = columns_twips[idx]
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.width = Mm(twips_to_mm(width))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_shading(cell)


def fill_address_cell(cell, label: str, lines: list[str], align) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(1.5)
    paragraph.paragraph_format.line_spacing = Pt(12.4)
    if label:
        label_run = paragraph.add_run(label)
        label_run.bold = True
        set_run_font(label_run, SANS, 9.5)
    for line in lines:
        if paragraph.runs:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, SERIF, 10.3)


def load_address_block(value: str | None) -> dict[str, Any]:
    if not value:
        return {}
    path = Path(value)
    if path.exists():
        return json.loads(path.read_text(encoding="utf-8"))
    if value.strip().startswith("{"):
        return json.loads(value)
    raise FileNotFoundError(f"Address block JSON not found: {value}")


def first_content_paragraph(doc):
    return next((p for p in doc.paragraphs if p.text.strip()), doc.paragraphs[0] if doc.paragraphs else None)


def insert_table_before(doc, reference):
    table = doc.add_table(rows=1, cols=2)
    body = doc._body._body
    body.remove(table._tbl)
    if reference is not None:
        reference._p.addprevious(table._tbl)
    else:
        body.append(table._tbl)
    return table


def ensure_city_date(doc, table, city_date: str) -> None:
    if not city_date:
        return
    date_like = re.compile(r"^\s*.+,\s*\d{1,2}\s+\S+\s+\d{4}")
    for paragraph in doc.paragraphs[:8]:
        if paragraph.text.strip() == city_date or date_like.match(paragraph.text):
            clear_paragraph(paragraph)
            set_run_font(paragraph.add_run(city_date), SERIF, 10.5)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            return
    reference = first_content_paragraph(doc)
    paragraph = reference.insert_paragraph_before(city_date) if reference is not None else doc.add_paragraph(city_date)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in paragraph.runs:
        set_run_font(run, SERIF, 10.5)
    paragraph._p.addnext(table._tbl)


def apply_address_block(doc, raw_value: str | None) -> None:
    data = load_address_block(raw_value)
    columns = tuple(data.get("columns_twips", DEFAULT_ADDRESS_COLS_TWIPS))
    if len(columns) != 2:
        raise ValueError("address-block columns_twips must contain two widths")

    if doc.tables:
        table = doc.tables[0]
    else:
        table = insert_table_before(doc, first_content_paragraph(doc))
    if len(table.rows) > 1 and not "".join(cell.text.strip() for cell in table.rows[0].cells):
        remove_row(table.rows[0])
    if data.get("city_date"):
        ensure_city_date(doc, table, str(data["city_date"]))

    if table.rows and len(table.rows[0].cells) >= 2 and ("sender" in data or "recipient" in data):
        fill_address_cell(
            table.rows[0].cells[0],
            str(data.get("sender_label", "Sender:")),
            [str(x) for x in data.get("sender", [])],
            WD_ALIGN_PARAGRAPH.LEFT,
        )
        fill_address_cell(
            table.rows[0].cells[1],
            str(data.get("recipient_label", "Recipient:")),
            [str(x) for x in data.get("recipient", [])],
            WD_ALIGN_PARAGRAPH.RIGHT if data.get("recipient_align", "right") == "right" else WD_ALIGN_PARAGRAPH.LEFT,
        )

    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_address_table_geometry(table, (int(columns[0]), int(columns[1])))
    remove_table_borders(table)


def keep_headings_with_next(doc) -> None:
    for style in doc.styles:
        if style.type == WD_STYLE_TYPE.PARAGRAPH and style.name.startswith("Heading "):
            style.paragraph_format.keep_with_next = True
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading "):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True


def shrink_verbatim(doc) -> None:
    try:
        style = doc.styles["Verbatim Char"]
    except KeyError:
        return
    set_style_font(style, MONO, 8)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if getattr(run.style, "name", "") == "Verbatim Char":
                set_run_font(run, MONO, 8)


def parse_args(argv: list[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Post-process a Quarto/Pandoc DOCX.")
    parser.add_argument("docx", type=Path, help="DOCX file to update in place.")
    parser.add_argument(
        "--geometry",
        default="A4,top=20,bottom=20,left=24,right=24",
        help="Page geometry, for example: A4,top=20,bottom=20,left=24,right=24.",
    )
    parser.add_argument("--running-header", help="Running header text, applied from page 2.")
    parser.add_argument(
        "--footer-page",
        nargs="?",
        const="",
        help="Add a PAGE-field footer; optional text appears on the left.",
    )
    parser.add_argument("--doc-label", help="Document kind label inserted above the title.")
    parser.add_argument(
        "--address-block",
        nargs="?",
        const="",
        help="Normalize the first table as a borderless address block; optional JSON path or inline JSON fills it.",
    )
    parser.add_argument(
        "--keep-with-next",
        action="store_true",
        help="Keep heading paragraphs with the paragraph that follows.",
    )
    parser.add_argument(
        "--boxes-to-tables",
        action="store_true",
        help="Repack KeyBox/MethodBox paragraph groups as one-cell tables with reliable Word print frames.",
    )
    parser.add_argument("--theme", default="think-tank", choices=theme_names())
    parser.add_argument("--accent", help="Override the theme accent with a 6 digit hex colour.")
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    args = parse_args(sys.argv[1:] if argv is None else argv)
    if not args.docx.exists():
        raise FileNotFoundError(args.docx)

    theme = get_theme(args.theme, args.accent)
    doc = Document(str(args.docx))
    apply_geometry(doc, args.geometry)
    shrink_verbatim(doc)
    if args.doc_label:
        add_doc_label(doc, args.doc_label, args.theme, theme)
    if args.running_header:
        set_running_header(doc, args.running_header, args.theme, theme)
    if args.footer_page is not None:
        set_page_footer(doc, args.footer_page, theme)
    if args.address_block is not None:
        apply_address_block(doc, args.address_block)
    converted_boxes = 0
    if args.boxes_to_tables:
        converted_boxes = convert_boxes_to_tables(doc, args.theme, theme)
    if args.keep_with_next:
        keep_headings_with_next(doc)

    doc.save(str(args.docx))
    suffix = f" ({converted_boxes} box table{'s' if converted_boxes != 1 else ''})" if args.boxes_to_tables else ""
    print(f"DOCX postprocess complete: {args.docx}{suffix}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
