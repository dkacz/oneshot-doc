"""oneshot-doc QA gates.

Usage:
    python3 qa_gates.py --qmd doc.qmd --pdf doc.pdf --budget 4 \
        [--numbers NUMBERS.md] [--docx doc.docx] [--docx-budget 4] \
        [--parity-warn-only] [--outdir qa]

Gates (fail | warn | pass | skip):
  pages               PDF and DOCX page counts within budget
  metanarration       no self-reference, production-process prose, or defensive prose
  dashes_source       no em/en dashes in the qmd source
  dashes_pdf          no em/en dashes in extracted PDF text
  dashes_docx         no em/en dashes in docx text (if --docx)
  numbers             every contextual number in the PDF text is in NUMBERS.md
  bibliography        in-text citations and bibliography entries match
  fonts               all PDF fonts embedded
  figure_diacritics   Polish figure labels keep diacritics and the palette font has glyphs
  pageshots           page PNGs rendered for the visual critic (qa/pages/)
  docx_preview        DOCX converted to PDF and page PNGs for inspection (if --docx)
  docx_parity         PDF and DOCX-preview page counts and page starts match (if --docx)

Exit code 1 if any gate fails. Results are also written to <outdir>/gates.json.
"""

from __future__ import annotations

import argparse
import ast
from collections import Counter
import json
import re
import shutil
import sys
import unicodedata
from pathlib import Path

from toolpaths import resolved_tool_path, run_tool

for _stream in (sys.stdout, sys.stderr):
    try:
        _stream.reconfigure(encoding="utf-8", errors="replace")
    except (AttributeError, ValueError):
        pass

DASHES = {"—": "em-dash", "–": "en-dash"}
NUM_CONTEXT = re.compile(
    r"(%|percent|percentage|procent|punkt|pp\b|p\.p\.|bn|mld|mln|tys\.|zł|euro|EUR|USD|dollar|dolar)",
    re.I,
)
META_PATTERNS = [
    ("ta notatka", re.compile(r"\bta\s+notatka\b", re.I)),
    ("tej notatki", re.compile(r"\btej\s+notatki\b", re.I)),
    ("niniejsz*", re.compile(r"\bniniejsz\w*\b", re.I)),
    ("ten dokument", re.compile(r"\bten\s+dokument\b", re.I)),
    ("tego dokumentu", re.compile(r"\btego\s+dokumentu\b", re.I)),
    ("ten brief", re.compile(r"\bten\s+brief\b", re.I)),
    ("ten raport", re.compile(r"\bten\s+raport\b", re.I)),
    ("to pismo", re.compile(r"\bto\s+pismo\b", re.I)),
    ("w tym dokumencie", re.compile(r"\bw\s+tym\s+dokumencie\b", re.I)),
    ("this report", re.compile(r"\bthis\s+report\b", re.I)),
    ("this document", re.compile(r"\bthis\s+document\b", re.I)),
    ("this brief", re.compile(r"\bthis\s+brief\b", re.I)),
    ("this memo", re.compile(r"\bthis\s+memo\b", re.I)),
    ("the present document", re.compile(r"\bthe\s+present\s+document\b", re.I)),
    ("brief produkcyjny", re.compile(r"\bbrief\w*\s+produkcyjn\w*\b", re.I)),
    ("kontrakt produkcyjny", re.compile(r"\bkontrakt\w*\s+produkcyjn\w*\b", re.I)),
    ("zgodnie z kontraktem", re.compile(r"\bzgodnie\s+z\s+kontraktem\b", re.I)),
    ("ustalenia z wywiadu", re.compile(r"\bustalen\w*\s+z\s+wywiadu\b", re.I)),
    ("rundy poprawek", re.compile(r"\brund\w*\s+poprawek\b", re.I)),
    ("production brief", re.compile(r"\bproduction\s+brief\b", re.I)),
    ("review round", re.compile(r"\breview\s+rounds?\b", re.I)),
    ("interview settlement", re.compile(r"\binterview\s+settlements?\b", re.I)),
    ("according to the contract", re.compile(r"\baccording\s+to\s+the\s+contract\b", re.I)),
    ("contract requires", re.compile(r"\bthe\s+contract\s+requires\b", re.I)),
]
DEFENSIVE_SUBJECT = re.compile(
    r"\b(?:dokument|notatka|brief|raport|pismo|document|memo|report)\b"
    r"[^.\n;:!?]{0,100}\b(?:nie\s+zawiera|nie\s+przedstawia|nie\s+zakłada|nie\s+zaklada|"
    r"nie\s+uruchamia|nie\s+rozstrzyga|does\s+not\s+contain|does\s+not\s+present|"
    r"does\s+not\s+assume|does\s+not\s+launch|does\s+not\s+resolve|does\s+not\s+decide)\b",
    re.I,
)
POLISH_ASCII_LABELS = [
    ("okolo", "około"),
    ("zl", "zł"),
    ("wedlug", "według"),
    ("zrodlo", "źródło"),
    ("wzgledem", "względem"),
    ("laczny", "łączny"),
    ("wiecej", "więcej"),
    ("srednio", "średnio"),
    ("wzrost gosp", "wzrost gospodarczy"),
]
POLISH_GLYPHS = "ąćęłńóśźż"
SKILL_ROOT = Path(__file__).resolve().parents[1]
ASSETS_DIR = SKILL_ROOT / "assets"


def run_resolved(name: str, args: list[str], **kw):
    kw.setdefault("capture_output", True)
    kw.setdefault("text", True)
    kw.setdefault("encoding", "utf-8")
    kw.setdefault("errors", "replace")
    return run_tool(name, args, **kw)


def pdf_text(pdf: Path) -> str:
    r = run_resolved("pdftotext", ["-layout", str(pdf), "-"])
    return r.stdout


def pdf_pages(pdf: Path) -> int:
    import pypdf

    return len(pypdf.PdfReader(str(pdf)).pages)


def pdf_page_texts(pdf: Path) -> list[str]:
    pages = []
    for page_no in range(1, pdf_pages(pdf) + 1):
        r = run_resolved("pdftotext", ["-f", str(page_no), "-l", str(page_no), "-layout", str(pdf), "-"])
        pages.append(r.stdout)
    return pages


def normalized_words(text: str) -> list[str]:
    text = unicodedata.normalize("NFKC", text.casefold()).replace("\u00ad", "")
    return re.findall(r"[\wąćęłńóśźż]+", text, re.UNICODE)


def normalized_line(line: str) -> str:
    return " ".join(normalized_words(line))


def normalized_running_line_key(line: str) -> str:
    return " ".join(word for word in normalized_words(line) if not word.isdigit())


def page_words_without_running_matter(page_texts: list[str]) -> list[list[str]]:
    line_counts: Counter[str] = Counter()
    for text in page_texts:
        seen_on_page = set()
        for line in text.splitlines():
            key = normalized_running_line_key(line)
            if key:
                seen_on_page.add(key)
        line_counts.update(seen_on_page)
    repeated = {line for line, count in line_counts.items() if count >= 2}

    pages = []
    for text in page_texts:
        kept_lines = []
        for line in text.splitlines():
            norm = normalized_line(line)
            key = normalized_running_line_key(line)
            if not norm or norm.isdigit() or key in repeated:
                continue
            kept_lines.append(line)
        pages.append(normalized_words("\n".join(kept_lines)))
    return pages


def seq_index(haystack: list[str], needle: list[str]) -> int:
    if not needle:
        return 0
    end = len(haystack) - len(needle) + 1
    for idx in range(max(0, end)):
        if haystack[idx: idx + len(needle)] == needle:
            return idx
    return -1


def page_start_matches(pdf_words: list[str], docx_words: list[str], size: int = 8, max_skip: int = 28) -> bool:
    if not pdf_words and not docx_words:
        return True
    if not pdf_words or not docx_words:
        return False
    sig_len = min(size, len(pdf_words), len(docx_words))
    if sig_len == 0:
        return False
    pdf_sig = pdf_words[:sig_len]
    docx_sig = docx_words[:sig_len]
    if pdf_sig == docx_sig:
        return True
    pdf_window = pdf_words[: max_skip + sig_len]
    docx_window = docx_words[: max_skip + sig_len]
    return seq_index(docx_window, pdf_sig) >= 0 or seq_index(pdf_window, docx_sig) >= 0


def signature(words: list[str], size: int = 8) -> str:
    return " ".join(words[:size]) if words else "(no page text)"


def docx_parity_gate(pdf: Path, docx_pdf: Path, warn_only: bool) -> dict:
    problems = []
    pdf_count = pdf_pages(pdf)
    docx_count = pdf_pages(docx_pdf)
    if pdf_count != docx_count:
        problems.append(f"page count differs: PDF has {pdf_count}, DOCX preview has {docx_count}")

    pdf_page_words = page_words_without_running_matter(pdf_page_texts(pdf))
    docx_page_words = page_words_without_running_matter(pdf_page_texts(docx_pdf))
    for idx in range(min(len(pdf_page_words), len(docx_page_words))):
        if page_start_matches(pdf_page_words[idx], docx_page_words[idx]):
            continue
        problems.append(
            "page "
            f"{idx + 1} starts differ after header/footer normalization: "
            f"PDF starts {signature(pdf_page_words[idx])!r}; "
            f"DOCX starts {signature(docx_page_words[idx])!r}"
        )

    status = "pass" if not problems else ("warn" if warn_only else "fail")
    return {"status": status, "detail": problems[:30] if problems else "PDF and DOCX page starts match"}


def docx_text(docx: Path) -> str:
    import docx as docxlib

    d = docxlib.Document(str(docx))
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                parts.append(c.text)
    return "\n".join(parts)


def convert_docx_to_pdf(docx: Path, outdir: Path) -> tuple[Path | None, str]:
    preview_dir = outdir / "docx_preview"
    if preview_dir.exists():
        shutil.rmtree(preview_dir)
    preview_dir.mkdir(parents=True, exist_ok=True)
    if resolved_tool_path("soffice") is None:
        return None, "soffice not found"
    profile_dir = preview_dir / "lo_profile"
    profile_dir.mkdir(parents=True, exist_ok=True)
    r = run_resolved(
        "soffice",
        [
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            f"-env:UserInstallation={profile_dir.resolve().as_uri()}",
            "--convert-to",
            "pdf",
            "--outdir",
            str(preview_dir),
            str(docx),
        ],
    )
    expected = preview_dir / (docx.stem + ".pdf")
    if expected.exists():
        return expected, str(expected)
    candidates = sorted(preview_dir.glob("*.pdf"))
    if candidates:
        return candidates[0], str(candidates[0])
    detail = (r.stderr or r.stdout or "DOCX conversion produced no PDF").strip()
    return None, detail[-500:]


def find_dashes(text: str, label: str):
    hits = []
    for i, line in enumerate(text.splitlines(), 1):
        for ch, name in DASHES.items():
            if ch in line:
                hits.append(f"{label}:{i}: {name}: {line.strip()[:90]}")
    return hits


def norm_num(tok: str) -> str:
    return tok.replace(" ", "").replace(" ", "").replace(",", ".").strip(".")


def ledger_numbers(numbers_md: Path):
    """All numeric tokens appearing anywhere in NUMBERS.md count as licensed."""
    text = numbers_md.read_text(encoding="utf-8")
    toks = re.findall(r"\d+(?:[.,]\d+)?", text)
    return {norm_num(t) for t in toks}


def numbers_gate(pdftext: str, licensed: set):
    """Flag unlicensed numbers in the body text, with the bibliography excluded.

    Decimals always need a licence. Bare integers need one only when a unit
    keyword follows immediately, so figure/section references and years stay quiet.
    """
    cut = re.search(r"(Przywołana literatura|Bibliografia|References|Bibliography)", pdftext)
    body = pdftext[: cut.start()] if cut else pdftext
    failures, seen = [], set()
    for m in re.finditer(r"\d+(?:[.,]\d+)?", body):
        tok = m.group(0)
        n = norm_num(tok)
        if n in licensed or n in seen:
            continue
        if tok.isdigit() and 1900 <= int(tok) <= 2100:
            continue
        has_decimal = ("," in tok) or ("." in tok)
        after = body[m.end(): m.end() + 18].replace("\n", " ")
        if has_decimal or (tok.isdigit() and NUM_CONTEXT.search(after)):
            ctx = body[max(0, m.start() - 35): m.end() + 35].replace("\n", " ")
            seen.add(n)
            failures.append(f"number {tok!r} is not licensed in NUMBERS.md | ...{ctx.strip()}...")
    return failures


def bib_gate(qmd_text: str):
    """Match in-text (Name YYYY) citations against the bibliography section."""
    m = re.search(
        r"^#+\s*(Przywołana literatura|Bibliografia|Literatura|References|Bibliography)\s*$",
        qmd_text,
        re.M | re.I,
    )
    if not m:
        m = re.search(r"(Przywołana literatura|Bibliografia|References|Bibliography)", qmd_text)
    if not m:
        return None, []
    body, bib = qmd_text[: m.start()], qmd_text[m.end():]
    bib_plain = unicodedata.normalize("NFC", bib)
    problems = []
    cited = set()
    for cm in re.finditer(r"\(([^()]{0,120}?(?:19|20)\d{2}[a-z]?)\)", body):
        for part in cm.group(1).split(";"):
            part = part.strip()
            ym = re.search(r"(19|20)\d{2}", part)
            if not ym:
                continue
            year = ym.group(0)
            name_m = re.search(r"([A-ZŁŚŻŹĆÓ][\wà-ſ-]+)", part)
            if not name_m:
                continue
            cited.add((name_m.group(1), year))
    for name, year in sorted(cited):
        if not re.search(re.escape(name) + r"[^\n]{0,160}" + year, bib_plain) and not re.search(
            re.escape(name) + r".{0,400}?" + year, bib_plain, re.S
        ):
            problems.append(f"citation ({name} {year}) has no bibliography entry")
    for em in re.finditer(r"^([A-ZŁŚŻŹĆÓ][\wà-ſ-]+)[^\n]{0,200}?((?:19|20)\d{2})", bib_plain, re.M):
        name, year = em.group(1), em.group(2)
        if not re.search(re.escape(name) + r"[^()]{0,80}" + year, body) and not re.search(
            re.escape(name) + r"\s+(i in\.|et al\.|i\s|and\s)?[^()]{0,80}" + year, body
        ):
            problems.append(f"bibliography entry {name} {year} is not cited in text (warning)")
    return True, problems


def fonts_gate(pdf: Path):
    r = run_resolved("pdffonts", [str(pdf)])
    bad = []
    for line in r.stdout.splitlines()[2:]:
        cols = line.split()
        if len(cols) >= 5 and "no" in cols[-4:-2]:
            bad.append(line.strip())
    return bad


def qmd_text_lines_outside_code(qmd_text: str) -> list[tuple[int, str]]:
    lines = []
    in_fence = False
    fence_re = re.compile(r"^\s*(```|~~~)")
    for no, line in enumerate(qmd_text.splitlines(), 1):
        if fence_re.match(line):
            in_fence = not in_fence
            continue
        if not in_fence:
            lines.append((no, line))
    return lines


def allowed_hit(quote: str, allows: list[str]) -> bool:
    norm_quote = unicodedata.normalize("NFC", quote).casefold()
    return any(unicodedata.normalize("NFC", item).casefold() in norm_quote for item in allows)


def metanarration_gate(qmd_text: str, allows: list[str]):
    hits = []
    for no, line in qmd_text_lines_outside_code(qmd_text):
        quote = " ".join(line.strip().split())
        if not quote or allowed_hit(quote, allows):
            continue
        for label, pattern in META_PATTERNS:
            if pattern.search(line):
                hits.append(f"line {no}: {label}: {quote[:180]}")
        if DEFENSIVE_SUBJECT.search(line):
            hits.append(f"line {no}: defensive document-subject sentence: {quote[:180]}")
    return hits


def string_literals(py_path: Path) -> list[tuple[int, str]]:
    text = py_path.read_text(encoding="utf-8")
    try:
        tree = ast.parse(text, filename=str(py_path))
    except SyntaxError:
        fallback = []
        for no, line in enumerate(text.splitlines(), 1):
            for match in re.finditer(r"(['\"])(.*?)(?<!\\)\1", line):
                fallback.append((no, match.group(2)))
        return fallback

    out: list[tuple[int, str]] = []
    for node in ast.walk(tree):
        if isinstance(node, ast.Constant) and isinstance(node.value, str):
            out.append((getattr(node, "lineno", 1), node.value))
        elif isinstance(node, ast.JoinedStr):
            parts = [part.value for part in node.values if isinstance(part, ast.Constant) and isinstance(part.value, str)]
            if parts:
                out.append((getattr(node, "lineno", 1), "".join(parts)))
    return out


def ascii_label_pattern(label: str) -> re.Pattern:
    if " " in label:
        return re.compile(re.escape(label), re.I)
    letters = "A-Za-zĄĆĘŁŃÓŚŹŻąćęłńóśźż"
    return re.compile(rf"(?<![{letters}]){re.escape(label)}(?![{letters}])", re.I)


def palette_font_glyphs_gate():
    sys.path.insert(0, str(ASSETS_DIR))
    try:
        from matplotlib import font_manager
        from matplotlib import ft2font
        import matplotlib.pyplot as plt
        import palette

        palette.setup_fonts()
        family = plt.rcParams.get("font.family", "Source Sans 3")
        if isinstance(family, list):
            family = family[0]
        font_path = font_manager.findfont(font_manager.FontProperties(family=family), fallback_to_default=False)
        cmap = ft2font.FT2Font(font_path).get_charmap()
        missing = [ch for ch in POLISH_GLYPHS if ord(ch) not in cmap]
        if missing:
            return [f"palette font {family!r} at {font_path} lacks glyphs: {''.join(missing)}"]
        return []
    except Exception as exc:  # pragma: no cover - depends on local font stack
        return [f"could not verify palette font glyphs: {exc}"]


def figure_diacritics_gate(qmd: Path):
    problems = []
    figures = qmd.parent / "figures"
    pngs = sorted(figures.glob("*.png")) if figures.exists() else []
    for png in pngs:
        source = png.with_suffix(".py")
        if not source.exists():
            problems.append(f"{png}: no matching source script {source.name} (warning)")
            continue
        for no, value in string_literals(source):
            for bad, good in POLISH_ASCII_LABELS:
                if ascii_label_pattern(bad).search(value):
                    quote = " ".join(value.split())
                    problems.append(f"{source}:{no}: ASCII Polish label {bad!r}; use {good!r} | {quote[:160]}")
    problems.extend(palette_font_glyphs_gate())
    return problems


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--qmd", required=True, type=Path)
    ap.add_argument("--pdf", required=True, type=Path)
    ap.add_argument("--docx", type=Path)
    ap.add_argument("--numbers", type=Path)
    ap.add_argument("--budget", required=True, type=int)
    ap.add_argument("--docx-budget", type=int, help="DOCX page budget; defaults to --budget.")
    ap.add_argument(
        "--parity-warn-only",
        action="store_true",
        help="Downgrade DOCX/PDF page-count and page-start parity failures to warnings.",
    )
    ap.add_argument("--meta-allow", action="append", default=[], help="Allow one metanarration phrase, repeatable.")
    ap.add_argument("--skip-figure-diacritics", action="store_true", help="Skip figure-label diacritics checks.")
    ap.add_argument("--outdir", type=Path, default=Path("qa"))
    a = ap.parse_args()
    a.outdir.mkdir(parents=True, exist_ok=True)
    gates = {}

    docx_pdf = None
    docx_convert_detail = ""
    if a.docx and a.docx.exists():
        docx_pdf, docx_convert_detail = convert_docx_to_pdf(a.docx, a.outdir)

    pdf_npages = pdf_pages(a.pdf)
    page_detail = [f"PDF: {pdf_npages} pages with budget {a.budget}"]
    page_failed = pdf_npages > a.budget
    if a.docx and a.docx.exists():
        if docx_pdf and docx_pdf.exists():
            docx_budget = a.docx_budget if a.docx_budget is not None else a.budget
            docx_npages = pdf_pages(docx_pdf)
            page_detail.append(f"DOCX: {docx_npages} pages with budget {docx_budget}")
            page_failed = page_failed or (docx_npages > docx_budget)
        else:
            page_detail.append(f"DOCX: conversion failed: {docx_convert_detail}")
            page_failed = True
    gates["pages"] = {"status": "fail" if page_failed else "pass", "detail": page_detail}

    qmd_text = a.qmd.read_text(encoding="utf-8")
    hits = metanarration_gate(qmd_text, a.meta_allow)
    gates["metanarration"] = {"status": "fail" if hits else "pass", "detail": hits[:40]}

    hits = find_dashes(qmd_text, a.qmd.name)
    gates["dashes_source"] = {"status": "fail" if hits else "pass", "detail": hits[:20]}

    ptext = pdf_text(a.pdf)
    hits = find_dashes(ptext, a.pdf.name)
    gates["dashes_pdf"] = {"status": "fail" if hits else "pass", "detail": hits[:20]}

    if a.docx and a.docx.exists():
        dtext = docx_text(a.docx)
        hits = find_dashes(dtext, a.docx.name)
        gates["dashes_docx"] = {"status": "fail" if hits else "pass", "detail": hits[:20]}

    if a.numbers and a.numbers.exists():
        fails = numbers_gate(ptext, ledger_numbers(a.numbers))
        gates["numbers"] = {"status": "fail" if fails else "pass", "detail": fails[:30]}
    else:
        gates["numbers"] = {"status": "skip", "detail": "NUMBERS.md is missing (required in production)"}

    res, problems = bib_gate(qmd_text)
    if res is None:
        gates["bibliography"] = {"status": "skip", "detail": "no bibliography section"}
    else:
        hard = [p for p in problems if "warning" not in p]
        status = "fail" if hard else ("warn" if problems else "pass")
        gates["bibliography"] = {"status": status, "detail": problems[:20]}

    bad = fonts_gate(a.pdf)
    gates["fonts"] = {"status": "fail" if bad else "pass", "detail": bad[:10]}

    if a.skip_figure_diacritics:
        gates["figure_diacritics"] = {"status": "skip", "detail": "skipped by --skip-figure-diacritics"}
    else:
        problems = figure_diacritics_gate(a.qmd)
        hard = [p for p in problems if "(warning)" not in p]
        status = "fail" if hard else ("warn" if problems else "pass")
        gates["figure_diacritics"] = {"status": status, "detail": problems[:30]}

    pages_dir = a.outdir / "pages"
    if pages_dir.exists():
        shutil.rmtree(pages_dir)
    pages_dir.mkdir(parents=True)
    run_resolved("pdftoppm", ["-png", "-r", "110", str(a.pdf), str(pages_dir / "page")])
    shots = sorted(pages_dir.glob("page-*.png"))
    gates["pageshots"] = {
        "status": "pass" if shots else "fail",
        "detail": f"{len(shots)} screenshots in {pages_dir}",
    }

    if a.docx and a.docx.exists():
        docx_shot_detail = ""
        if docx_pdf and docx_pdf.exists():
            docx_pages_dir = a.outdir / "docx_pages"
            if docx_pages_dir.exists():
                shutil.rmtree(docx_pages_dir)
            docx_pages_dir.mkdir(parents=True)
            run_resolved("pdftoppm", ["-png", "-r", "110", str(docx_pdf), str(docx_pages_dir / "page")])
            docx_shots = sorted(docx_pages_dir.glob("page-*.png"))
            docx_shot_detail = f"; {len(docx_shots)} screenshots in {docx_pages_dir}"
        gates["docx_preview"] = {
            "status": "pass" if docx_pdf and docx_pdf.exists() else "fail",
            "detail": (str(docx_pdf) + docx_shot_detail) if docx_pdf and docx_pdf.exists() else docx_convert_detail,
        }
        if docx_pdf and docx_pdf.exists():
            gates["docx_parity"] = docx_parity_gate(a.pdf, docx_pdf, a.parity_warn_only)
        else:
            gates["docx_parity"] = {
                "status": "skip",
                "detail": "DOCX preview conversion failed; docx_preview gate reports the conversion error.",
            }
    else:
        gates["docx_parity"] = {"status": "skip", "detail": "no --docx supplied"}

    (a.outdir / "gates.json").write_text(json.dumps(gates, ensure_ascii=False, indent=2), encoding="utf-8")
    worst = "pass"
    for name, gate in gates.items():
        detail = gate["detail"]
        print(f"[{gate['status'].upper():4}] {name}: {detail if isinstance(detail, str) else ''}")
        for item in (detail if isinstance(detail, list) else []):
            print(f"        - {item}")
        if gate["status"] == "fail":
            worst = "fail"
    print(f"\nGate result: {worst.upper()} (details: {a.outdir / 'gates.json'})")
    sys.exit(1 if worst == "fail" else 0)


if __name__ == "__main__":
    main()
